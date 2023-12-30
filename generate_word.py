from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import re


def add_content(doc, content_text):
    content = doc.add_paragraph(content_text)
    content_run = content.runs[0]
    content_run.font.name = "Aptos"
    content_run.font.size = Pt(23)

def set_header(doc, title_text):
    section = doc.sections[0]
    header = section.header

    # Add a paragraph to the header containing the title
    title_paragraph = header.paragraphs[0]
    title_run = title_paragraph.add_run(title_text)
    title_run.font.name = "Microsoft JhengHei"
    title_run.font.size = Pt(28)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def generate_word_doc(path, title, content):
    # Create a new Word document
    doc = Document()
    add_content(doc, content_text)

    # Task 3: Set header with title
    set_header(doc, title_text)

    # Save the document
    doc.save(f"{path}/{clean_filename(title)}.docx")

def clean_filename(filename):
    # Define a regex pattern to match invalid characters
    invalid_chars_pattern = re.compile(r'[\\/:"*?<>|]+')

    # Replace invalid characters with underscores
    cleaned_filename = re.sub(invalid_chars_pattern, '_', filename)

    return cleaned_filename

if __name__ == "__main__":
    # Create a new Word document
    doc = Document()

    # Task 1: Add title
    title_text = "Title Line"

    # Task 2: Add content
    content_text = "This is the content line."
    add_content(doc, content_text)

    # Task 3: Set header with title
    set_header(doc, title_text)

    # Save the document
    doc.save("output.docx")
